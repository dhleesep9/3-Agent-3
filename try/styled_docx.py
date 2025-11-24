from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

def create_styled_report(report_text: str, output_path="analysis_report.docx"):
    doc = Document()

    # ============================
    # 기본 스타일 설정
    # ============================
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    # ============================
    # Title
    # ============================
    title = doc.add_paragraph()
    run = title.add_run("Negative Review Analysis Report")
    run.font.size = Pt(20)
    run.bold = True
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 여백
    doc.add_paragraph("")

    # ============================
    # 본문 파싱 & 스타일 적용
    # ============================
    lines = report_text.split("\n")

    for line in lines:
        line = line.strip()

        if not line:
            continue

        # -------------------------
        # 부제(## Subtitle)
        # -------------------------
        if line.startswith("## "):
            subtitle = doc.add_paragraph()
            subtitle_run = subtitle.add_run(line.replace("## ", ""))
            subtitle_run.font.size = Pt(14)
            subtitle_run.bold = True
            subtitle_run.font.name = '맑은 고딕'
            subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            continue

        # -------------------------
        # 일반 문단
        # -------------------------
        p = doc.add_paragraph()
        p_run = p.add_run(line)
        p_run.font.size = Pt(11)
        p_run.font.name = '맑은 고딕'
        p_run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

        # 줄간격
        p.paragraph_format.line_spacing = 1.35
        p.paragraph_format.space_after = Pt(10)

    # 저장
    doc.save(output_path)
    print(f"📄 DOCX 생성 완료: {output_path}")